"""DOCX (Microsoft Word) loader.

Extracts the body text of a .docx file. Each paragraph becomes a line in
the resulting text; sections / tables are flattened to plain text only.
"""
from __future__ import annotations

from pathlib import Path

import docx

from notebooklm.ingestion.loaders.base import Document


class DocxLoader:
    """Load a .docx file, return paragraphs as a single `Document`."""

    def load(self, path: str | Path) -> list[Document]:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(p)
        d = docx.Document(str(p))
        paragraphs = [para.text for para in d.paragraphs if para.text.strip()]
        if not paragraphs:
            return []
        text = "\n\n".join(paragraphs)
        return [Document(text=text, source=str(p), metadata={"path": str(p)})]
